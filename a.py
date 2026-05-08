import streamlit as st
import re, random
from io import BytesIO
from copy import deepcopy
from collections import defaultdict
from docx import Document
from docx.shared import Inches
from docx.table import _Cell
import zipfile

st.set_page_config(page_title="EndSem QB Generator (Stable Multi-Set)", layout="wide")

# -----------------------
# Regex helpers
# -----------------------
BLOOM_RE = re.compile(r'K\s*([1-6])', re.I)
CO_RE = re.compile(r'CO\s*[_:]?\s*(\d+)', re.I)
UNIT_RE = re.compile(r'Unit\s*[-:]?\s*(\d+)', re.I)
DIGIT_ONLY_RE = re.compile(r'^\s*(\d+)\s*$')
NUM_PREFIX_RE = re.compile(r'^\s*(\d+)\s*[\.\)]')

# -----------------------
# XML helpers
# -----------------------
def _copy_element(elem):
    return deepcopy(elem)

def replace_cell_with_cell(target_cell, src_cell):
    """Copy entire DOCX cell XML (text + equations) to bypass formatting loss"""
    t_tc = target_cell._tc
    s_tc = src_cell._tc
    for child in list(t_tc):
        t_tc.remove(child)
    for child in list(s_tc):
        t_tc.append(_copy_element(child))

def extract_images_from_cell(cell):
    """Extract images as binary blobs from the Word cell"""
    images = []
    for blip in cell._element.xpath(".//*[local-name()='blip']"):
        rId = blip.get(
            "{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed"
        )
        if rId:
            part = cell.part.related_parts[rId]
            images.append(BytesIO(part.blob))
    return images

# -----------------------
# Parse Question Bank
# -----------------------
def extract_questions_from_bank_docx(uploaded_file):
    doc = Document(BytesIO(uploaded_file.read()))
    questions = []
    qid = 0

    for table in doc.tables:
        for row in table.rows:
            try:
                cells_text = [c.text.strip() for c in row.cells]
            except ValueError:
                # Skip rows that cause internal python-docx errors due to complex merges
                continue

            if not any(cells_text):
                continue

            joined = " ".join(cells_text)
            bm = BLOOM_RE.search(joined)
            if not bm:
                continue

            # Improved Unit Detection: Check for "Unit X" or a standalone digit in cells
            unit_val = None
            um = UNIT_RE.search(joined)
            if um:
                unit_val = int(um.group(1))
            else:
                for txt in cells_text:
                    dm = DIGIT_ONLY_RE.match(txt)
                    if dm:
                        unit_val = int(dm.group(1))
                        break

            cm = CO_RE.search(joined)

            try:
                # The question is usually in the cell with the most text
                main_idx = max(range(len(row.cells)), key=lambda i: len(row.cells[i].text))
                main_cell = row.cells[main_idx]
            except ValueError:
                continue

            qid += 1
            questions.append({
                "id": qid,
                "unit": unit_val,
                "co": int(cm.group(1)) if cm else None,
                "bloom": int(bm.group(1)),
                "cell": main_cell, 
                "images": extract_images_from_cell(main_cell)
            })

    return questions

# -----------------------
# Parse Template (Merged Cell Safe)
# -----------------------
def parse_template_slots(template_file):
    doc = Document(template_file)
    slots = []

    for ti, table in enumerate(doc.tables):
        for ri, row in enumerate(table.rows):
            # Safe access to cells using low-level XML list to prevent "tc element" errors
            tr = row._tr
            for ci, tc in enumerate(tr.tc_lst):
                cell = _Cell(tc, table)
                txt = cell.text.strip()
                if not txt:
                    continue

                if txt.upper() in ["OR", "(OR)", "( OR )"]:
                    slots.append({
                        "table_index": ti,
                        "row_index": ri,
                        "cell_index": ci,
                        "slot_num": None,
                        "is_or": True
                    })
                    continue

                m = NUM_PREFIX_RE.match(txt)
                if m:
                    slots.append({
                        "table_index": ti,
                        "row_index": ri,
                        "cell_index": ci,
                        "slot_num": int(m.group(1)),
                        "is_or": False
                    })
    return slots

# -----------------------
# Logic helpers
# -----------------------
def allowed_blooms_for_slot_num(slot_num):
    if 1 <= slot_num <= 10:
        return [1, 2, 3]
    if 11 <= slot_num <= 20:
        return [4, 5]
    if 21 <= slot_num <= 22:
        return [6]
    return [1, 2, 3, 4, 5, 6]

def allowed_unit_for_slot_num(slot_num, part_c_unit):
    # Rule: q1,2->U1, q3,4->U2, q5,6->U3, q7,8->U4, q9,10->U5
    # Same pattern repeats for q11-q20
    if slot_num in [1, 2, 11, 12]: return 1
    if slot_num in [3, 4, 13, 14]: return 2
    if slot_num in [5, 6, 15, 16]: return 3
    if slot_num in [7, 8, 17, 18]: return 4
    if slot_num in [9, 10, 19, 20]: return 5
    if slot_num in [21, 22]: return part_c_unit
    return None

def find_or_pairs(slots):
    pairs = []
    for i, s in enumerate(slots):
        if s["is_or"]:
            left = right = None
            for j in range(i-1, -1, -1):
                if slots[j].get("slot_num"):
                    left = slots[j]
                    break
            for k in range(i+1, len(slots)):
                if slots[k].get("slot_num"):
                    right = slots[k]
                    break
            if left and right:
                pairs.append((left, right))
    return pairs

# -----------------------
# Select Questions
# -----------------------
def select_questions(entries, or_pairs, questions, part_c_unit):
    selected = {}
    used_q_ids = set() 

    by_unit_bloom = defaultdict(lambda: defaultdict(list))
    for q in questions:
        if q["unit"] is not None:
            by_unit_bloom[q["unit"]][q["bloom"]].append(q)

    def get_unique_question(pool):
        available = [q for q in pool if q["id"] not in used_q_ids]
        if not available:
            # Fallback: repeat a question if bank is exhausted
            chosen_q = random.choice(pool) if pool else None
        else:
            chosen_q = random.choice(available)
        
        if chosen_q:
            used_q_ids.add(chosen_q["id"])
        return chosen_q

    def pick_for_slot(slot_num):
        target_unit = allowed_unit_for_slot_num(slot_num, part_c_unit)
        target_blooms = allowed_blooms_for_slot_num(slot_num)
        
        pool = []
        for b in target_blooms:
            pool.extend(by_unit_bloom.get(target_unit, {}).get(b, []))
        
        if not pool:
            return "UNIT_NOT_FOUND"
        return get_unique_question(pool)

    # Process OR pairs first to ensure choices are matched properly
    for left, right in or_pairs:
        q_l = pick_for_slot(left["slot_num"])
        if q_l: selected[(left["table_index"], left["row_index"], left["cell_index"])] = q_l
        
        q_r = pick_for_slot(right["slot_num"])
        if q_r: selected[(right["table_index"], right["row_index"], right["cell_index"])] = q_r
            
    # Process remaining single slots
    for e in entries:
        coord = (e["table_index"], e["row_index"], e["cell_index"])
        if coord in selected: continue
        
        q_val = pick_for_slot(e["slot_num"])
        if q_val: selected[coord] = q_val

    return selected, used_q_ids

# -----------------------
# Assemble DOCX (Merged Cell Safe)
# -----------------------
def assemble_doc(template_file, selected_map):
    doc = Document(template_file)
    
    for (ti, ri, ci), q in selected_map.items():
        table = doc.tables[ti]
        row = table.rows[ri]
        tr = row._tr
        cells_xml = tr.tc_lst
        
        # Ensure row has enough columns (Number, Question, CO, Bloom)
        if len(cells_xml) < 4: continue

        q_cell = _Cell(cells_xml[1], table)
        co_cell = _Cell(cells_xml[2], table)
        k_cell = _Cell(cells_xml[3], table)

        q_cell.text = ""
        co_cell.text = ""
        k_cell.text = ""

        if q == "UNIT_NOT_FOUND":
            q_cell.text = "unit not in given qb"
        else:
            replace_cell_with_cell(q_cell, q["cell"])
            # Remove existing drawing XML from copied cell to avoid duplicates/errors
            for p in q_cell.paragraphs:
                for run in list(p.runs):
                    drawing = run._element.find('.//w:drawing', namespaces={'w': "http://schemas.openxmlformats.org/wordprocessingml/2006/main"})
                    if drawing is not None:
                        p._element.remove(run._element)
            
            # Explicitly re-insert images found in the question bank cell
            for img in q.get("images", []):
                img.seek(0)
                q_cell.add_paragraph().add_run().add_picture(img, width=Inches(2.5))
            
            co_cell.add_paragraph(f"CO{q.get('co','')}")
            k_cell.add_paragraph(f"K{q.get('bloom','')}")

    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf

# -----------------------
# Streamlit UI
# -----------------------
st.title("📘 End-Sem Question Paper Generator (Stable Multi-Set)")

with st.sidebar:
    st.header("Settings")
    template_file = st.file_uploader("Upload Template DOCX", type="docx")
    bank_file = st.file_uploader("Upload Question Bank DOCX", type="docx")
    n_sets = st.number_input("Number of Sets", 1, 10, 2)
    part_c_unit = st.selectbox("Select Unit for Part C (q21, q22)", [1, 2, 3, 4, 5], index=4)

if st.button("Generate Question Papers"):
    if not template_file or not bank_file:
        st.error("Upload both template and question bank files.")
        st.stop()
        
    questions = extract_questions_from_bank_docx(bank_file)
    st.success(f"✅ {len(questions)} questions loaded from bank.")

    slots = parse_template_slots(template_file)
    # Filter entries to only include those with slot numbers
    entries = [s for s in slots if s.get("slot_num")]
    or_pairs = find_or_pairs(slots)
    
    if not entries:
        st.warning("No question slots (e.g., '1.', '2.') detected in the template.")
        st.stop()
        
    buffers = {}
    all_used_q_ids = []
    
    for i in range(n_sets):
        selected, used_q_ids_in_set = select_questions(entries, or_pairs, questions, part_c_unit)
        all_used_q_ids.extend([qid for qid in used_q_ids_in_set if isinstance(qid, int)])

        template_file.seek(0) 
        buf = assemble_doc(template_file, selected)
        buffers[f"Set_{i+1}.docx"] = buf

        st.download_button(
            f"📄 Download Set {i+1}",
            buf,
            file_name=f"Set_{i+1}.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )
        
    # --- Inter-set analysis ---
    if n_sets >= 2 and all_used_q_ids:
        total_questions_generated = len(all_used_q_ids)
        unique_questions_used = len(set(all_used_q_ids))
        total_repetitions = total_questions_generated - unique_questions_used
        repetition_percentage = (total_repetitions / total_questions_generated) * 100 if total_questions_generated > 0 else 0
        
        st.markdown("---")
        st.header("📊 Repetition Analysis")
        st.info(f"Inter-Set Repetition Percentage: {repetition_percentage:.2f}%")

    # --- ZIP Packaging ---
    zip_buf = BytesIO()
    with zipfile.ZipFile(zip_buf, "w") as z:
        for name, buf in buffers.items():
            z.writestr(name, buf.getvalue())
    zip_buf.seek(0)
    st.download_button("🗂 Download All Sets (ZIP)", zip_buf, "All_Sets.zip")